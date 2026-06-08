"""Convierte INFORME-TECNICO.md a INFORME-TECNICO.docx"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "INFORME-TECNICO.md"
DOCX_PATH = ROOT / "INFORME-TECNICO.docx"


def ConfigurarEstilos(documento):
    estilo_normal = documento.styles["Normal"]
    estilo_normal.font.name = "Calibri"
    estilo_normal.font.size = Pt(11)

    for nivel, tamano in [(1, 22), (2, 16), (3, 13), (4, 12)]:
        estilo = documento.styles[f"Heading {nivel}"]
        estilo.font.name = "Calibri"
        estilo.font.size = Pt(tamano)
        estilo.font.bold = True
        estilo.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)


def AgregarBordeTabla(tabla):
    tbl = tabla._tbl
    tbl_pr = tbl.tblPr
    bordes = OxmlElement("w:tblBorders")
    for borde in ("top", "left", "bottom", "right", "insideH", "insideV"):
        nodo = OxmlElement(f"w:{borde}")
        nodo.set(qn("w:val"), "single")
        nodo.set(qn("w:sz"), "4")
        nodo.set(qn("w:color"), "CCCCCC")
        bordes.append(nodo)
    tbl_pr.append(bordes)


def AgregarTextoConFormato(parrafo, texto):
    partes = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", texto)
    for parte in partes:
        if not parte:
            continue
        if parte.startswith("**") and parte.endswith("**"):
            run = parrafo.add_run(parte[2:-2])
            run.bold = True
        elif parte.startswith("`") and parte.endswith("`"):
            run = parrafo.add_run(parte[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        else:
            parrafo.add_run(parte)


def EsLineaTabla(linea):
    return linea.strip().startswith("|") and linea.strip().endswith("|")


def EsSeparadorTabla(linea):
    return bool(re.match(r"^\|\s*[-: ]+\|\s*([-: ]+\|\s*)+$", linea.strip()))


def ParsearCeldas(linea):
    return [celda.strip() for celda in linea.strip().strip("|").split("|")]


def AgregarTabla(documento, filas):
    if len(filas) < 1:
        return
    columnas = len(filas[0])
    tabla = documento.add_table(rows=len(filas), cols=columnas)
    tabla.style = "Table Grid"
    AgregarBordeTabla(tabla)

    for i, fila in enumerate(filas):
        for j, celda in enumerate(fila):
            celda_doc = tabla.rows[i].cells[j]
            celda_doc.text = ""
            p = celda_doc.paragraphs[0]
            AgregarTextoConFormato(p, celda)
            if i == 0:
                for run in p.runs:
                    run.bold = True

    documento.add_paragraph()


def AgregarBloqueCodigo(documento, lineas):
    p = documento.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("\n".join(lineas))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)


def ConvertirMarkdownADocx(ruta_md, ruta_docx):
    contenido = ruta_md.read_text(encoding="utf-8")
    lineas = contenido.splitlines()

    documento = Document()
    ConfigurarEstilos(documento)

    titulo = documento.add_heading("Informe Técnico — CloudPix Studio", 0)
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    i = 0
    en_codigo = False
    bloque_codigo = []
    filas_tabla = []

    while i < len(lineas):
        linea = lineas[i]

        if linea.strip().startswith("```"):
            if en_codigo:
                AgregarBloqueCodigo(documento, bloque_codigo)
                bloque_codigo = []
                en_codigo = False
            else:
                en_codigo = True
            i += 1
            continue

        if en_codigo:
            bloque_codigo.append(linea)
            i += 1
            continue

        if EsLineaTabla(linea):
            if not EsSeparadorTabla(linea):
                filas_tabla.append(ParsearCeldas(linea))
            i += 1
            continue

        if filas_tabla:
            AgregarTabla(documento, filas_tabla)
            filas_tabla = []

        if linea.strip() == "---":
            documento.add_paragraph("_" * 72)
            i += 1
            continue

        if linea.startswith("# "):
            documento.add_heading(linea[2:].strip(), level=1)
            i += 1
            continue

        if linea.startswith("## "):
            documento.add_heading(linea[3:].strip(), level=2)
            i += 1
            continue

        if linea.startswith("### "):
            documento.add_heading(linea[4:].strip(), level=3)
            i += 1
            continue

        if linea.startswith("#### "):
            documento.add_heading(linea[5:].strip(), level=4)
            i += 1
            continue

        if linea.startswith("> "):
            p = documento.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            AgregarTextoConFormato(p, linea[2:].strip())
            i += 1
            continue

        if re.match(r"^\d+\.\s", linea.strip()):
            p = documento.add_paragraph(style="List Number")
            AgregarTextoConFormato(p, re.sub(r"^\d+\.\s", "", linea.strip()))
            i += 1
            continue

        if linea.strip().startswith("- [ ] ") or linea.strip().startswith("- [x] "):
            marcado = linea.strip().startswith("- [x] ")
            texto = linea.strip()[6:]
            p = documento.add_paragraph(style="List Bullet")
            prefijo = "☑ " if marcado else "☐ "
            AgregarTextoConFormato(p, prefijo + texto)
            i += 1
            continue

        if linea.strip().startswith("- "):
            p = documento.add_paragraph(style="List Bullet")
            AgregarTextoConFormato(p, linea.strip()[2:])
            i += 1
            continue

        if linea.strip() == "":
            i += 1
            continue

        p = documento.add_paragraph()
        AgregarTextoConFormato(p, linea.strip())
        i += 1

    if filas_tabla:
        AgregarTabla(documento, filas_tabla)

    pie = documento.add_paragraph()
    pie.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = pie.add_run("CloudPix Studio — Informe técnico generado automáticamente")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    documento.save(ruta_docx)
    return ruta_docx


if __name__ == "__main__":
    if not MD_PATH.exists():
        print(f"No se encontró: {MD_PATH}", file=sys.stderr)
        sys.exit(1)

    salida = ConvertirMarkdownADocx(MD_PATH, DOCX_PATH)
    print(f"Generado: {salida}")
